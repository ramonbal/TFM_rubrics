import os
import smtplib
from email.message import EmailMessage
import xlwings as xw
from docx import Document
import re

# Constants for configuration
GENERAR_PDF  = True  # Set to False to skip PDF generation
ENVIAR_EMAIL = True  # Set to False to skip email sending

# Define file paths and names. These should be relative to the script's directory
fitxer_dades                   = "committees.xlsx"
plantilla_comite               = "EvaluationGuidelinesCommitteEN.xlsx"
fitxer_plantilla_email_signar  = "mail_instruccions_signar.docx"

# Columnes on es troben el nom i el títol (convertides a lletres)
col_estudiant = "A"  # Columna 1
col_titol     = "B"  # Columna 2

def llegeix_plantilla_email(fitxer_plantilla):
    """
    Read email template from Word document, extract plain text for parameter substitution,
    but keep the original document structure for formatting
    """
    from docx import Document
    
    doc = Document(fitxer_plantilla)
    
    # First, get the plain text version for parameter replacement
    text_content = []
    for paragraph in doc.paragraphs:
        text_content.append(paragraph.text)
    
    plain_text = '\n\n'.join(text_content)
    
    # Return both the plain text and the document object
    return {'text': plain_text, 'doc': doc}

def substitueix_parametres_amb_format(doc_info, parametres):
    """
    Replace parameters in Word document while preserving formatting and images
    """
    doc = doc_info['doc']
    plain_text = doc_info['text']
    
    # Check which parameters are available in the template
    replacements = {}
    for param, valor in parametres.items():
        pattern = f'[{param}]'
        if pattern in plain_text:
            replacements[pattern] = str(valor)
        else:
            print(f"⚠️ Paràmetre [{param}] no trobat a la plantilla")
    
    # Convert document to HTML while preserving formatting
    html_content = ['<html><body>']
    
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            html_content.append('<br>')
            continue
            
        # Build paragraph HTML
        para_html = '<p>'
        
        for run in paragraph.runs:
            # Handle images
            if run._element.xpath('.//a:blip'):
                para_html += process_image(run, doc)
            else:
                # Handle text with formatting
                para_html += process_text_run(run)
        
        para_html += '</p>'
        
        # Clean up fragmented formatting and apply parameter replacements
        para_html = clean_and_replace(para_html, replacements)
        html_content.append(para_html)
    
    html_content.append('</body></html>')
    return '\n'.join(html_content)

def process_image(run, doc):
    """Extract and encode image from a run"""
    try:
        blip_elements = run._element.xpath('.//a:blip')
        for blip in blip_elements:
            rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId:
                import base64
                image_part = doc.part.related_parts[rId]
                image_data = base64.b64encode(image_part.blob).decode()
                content_type = image_part.content_type
                return f'<img src="data:{content_type};base64,{image_data}" style="max-width:100%; height:auto;" />'
    except Exception as e:
        print(f"⚠️ Error processant imatge: {e}")
    return '[IMAGE]'

def process_text_run(run):
    """Process a text run with formatting"""
    text = run.text
    
    # Escape HTML-like text (like <filename.ext>)
    text = text.replace('<', '&lt;').replace('>', '&gt;')
    
    # Apply formatting
    if run.bold:
        text = f'<strong>{text}</strong>'
    if run.italic:
        text = f'<em>{text}</em>'
    if run.underline:
        text = f'<u>{text}</u>'
    
    return text

def clean_and_replace(html, replacements):
    """Clean fragmented formatting and apply parameter replacements"""
    # Clean up fragmented formatting tags
    html = html.replace('</strong><strong>', '')
    html = html.replace('</em><em>', '')
    html = html.replace('</u><u>', '')
    
    # Apply parameter replacements
    for pattern, valor in replacements.items():
        html = html.replace(pattern, valor)
    
    return html
    
def substitueix_parametres(plantilla, parametres):
    """
    Replace parameters in template. Parameters in [] are replaced with values from dict
    Simple text-based replacement
    """
    resultat = plantilla
    
    # Simple parameter replacement
    for param, valor in parametres.items():
        pattern = f'[{param}]'
        if pattern in resultat:
#            print(f"✓ Substituint {pattern} -> {valor}")
            resultat = resultat.replace(pattern, str(valor))
        else:
            print(f"⚠️ Paràmetre [{param}] no trobat a la plantilla")
    
#    print(f"🔍 Resultat final:")
#    print(resultat[:200] + "..." if len(resultat) > 200 else resultat)
    return resultat

def excel_to_pdf(excel_path, pdf_path=None):
    """
    Convert Excel file to PDF using xlwings
    """
    if pdf_path is None:
        pdf_path = os.path.splitext(excel_path)[0] + ".pdf"
    
    # Check if Excel file exists
    if not os.path.exists(excel_path):
        print(f"❌ El fitxer Excel no existeix: {excel_path}")
        return None
    
    # Delete existing PDF file if it exists (to avoid overwrite issues)
    if os.path.exists(pdf_path):
        try:
            os.remove(pdf_path)
            #print(f"🗑️ PDF existent eliminat: {os.path.basename(pdf_path)}")
        except Exception as e:
            print(f"⚠️ No s'ha pogut eliminar el PDF existent: {e}")
            # Try to continue anyway
    
    try:
        # Open Excel file, activate it, and convert to PDF
        wb = app.books.open(os.path.abspath(excel_path))
        wb.activate()
        wb.to_pdf(path=os.path.abspath(pdf_path))
        wb.close()
        
        # Verify PDF was created
        if os.path.exists(pdf_path):
            return pdf_path
        else:
            print(f"❌ PDF no s'ha creat: {pdf_path}")
            return None
            
    except Exception as e:
        print(f"❌ Error convertint a PDF: {e}")
        print(f"   Tipus d'error: {type(e).__name__}")
        return None

def envia_email(params):
    """
    Send email with PDF attachment using template from Word document
    """
    login = ["ramon@cvc.uab.cat", "Quanhisoc"]
    # login = ["ramon.baldrich@uab.cat", "Nomar1uab."]
    login = ["mcv@cvc.uab.cat", "Mcv-email"]
    
    # Define parameters for template substitution
    parametres = {
        'student_name': params['nom_estudiant'], 
        'role':         params['rol'], 
        'title':        params['titol'],
    }
    
    # Substitute parameters in template with formatting
    cos_email_html = substitueix_parametres_amb_format(params['plantilla_email'], parametres)

    msg = EmailMessage()
    msg["Subject"] = params['assumpte']
    msg["From"]    = login[0]
    msg["To"]      = params['destinatari']
    msg["CC"]      = "ramon@cvc.uab.cat"
    
    # Set both plain text and HTML content
    import re
    cos_email_text_clean = re.sub(r'<[^>]+>', '', cos_email_html)  # Remove HTML tags
    cos_email_text_clean = cos_email_text_clean.replace('&nbsp;', ' ')   # Replace HTML entities
    
    msg.set_content(cos_email_text_clean)                          # Plain text version
    msg.add_alternative(cos_email_html, subtype='html')      # HTML version

    # Add PDF attachment
    try:
        with open(params['fitxer_pdf'], "rb") as f:
            msg.add_attachment(
                f.read(), 
                maintype="application", 
                subtype="pdf", 
                filename=os.path.basename(params['fitxer_pdf'])
            )
        # print(f"✓ Fitxer adjunt: {os.path.basename(params['fitxer_pdf'])}")
    except Exception as e:
        print(f"❌ Error adjuntant fitxer: {e}")
        return False

    # Try different SMTP configurations for Microsoft Outlook services
    smtp_configs = [
        # Microsoft Office 365 / Outlook.com SMTP
        {"host": "smtp.office365.com", "port": 587, "use_tls": True, "use_ssl": False},
        {"host": "smtp-mail.outlook.com", "port": 587, "use_tls": True, "use_ssl": False},
        # University specific Microsoft Exchange
        {"host": "outlook.uab.cat", "port": 587, "use_tls": True, "use_ssl": False},
        {"host": "smtp.uab.cat", "port": 587, "use_tls": True, "use_ssl": False},
        # Alternative Microsoft configurations
        {"host": "smtp-mail.outlook.com", "port": 25, "use_tls": True, "use_ssl": False},
        {"host": "smtp.office365.com", "port": 25, "use_tls": True, "use_ssl": False},
        # Legacy configurations as fallback
        {"host": "cvc.uab.cat", "port": 587, "use_tls": True, "use_ssl": False},
        {"host": "cvc.uab.cat", "port": 25, "use_tls": False, "use_ssl": False},
    ]

    for i, config in enumerate(smtp_configs, 1):
        try:
            # print(f"Intent {i}: {config['host']}:{config['port']} (TLS:{config['use_tls']}, SSL:{config['use_ssl']})")            
            # Use regular SMTP
            with smtplib.SMTP(config['host'], config['port']) as smtp:
                if config['use_tls']:
                    smtp.starttls()
                smtp.login(login[0], login[1])
                smtp.send_message(msg)
                # print(f"✓ Email enviat correctament amb configuració {i}")
                return True
                    
        except Exception as e:
            print(f"❌ Configuració {i} fallida: {e}")
            continue
    
    print("❌ No s'ha pogut enviar l'email amb cap configuració")
    return False


# Canvia el directori de treball al directori del script
# Get the directory of the current script
script_dir = os.path.dirname(os.path.abspath(__file__))
os.chdir(script_dir)  # Change the current working directory to the script's directory

# Initialize xlwings app with better error handling
app = xw.App(visible=False)
# Disable alerts to prevent Excel prompts
app.display_alerts = False
# Disable screen updating for better performance
app.screen_updating = False

try:
    # Check if the data file exists
    if not os.path.exists(fitxer_dades):
        raise FileNotFoundError(f"El fitxer de dades '{fitxer_dades}' no existeix.")
    
    # Open the data file
    wb_dades = app.books.open(os.path.abspath(fitxer_dades))
    ws_dades = wb_dades.sheets.active
    
    # Get max row for processing using xlwings
    max_row = ws_dades.range('A1').expand('down').last_cell.row
    max_row = 2  # For testing purposes, set max_row to 4
    
    # Read email template
    plantilla_email_signar = llegeix_plantilla_email(fitxer_plantilla_email_signar)

    print(f"Processant {max_row - 1} files de dades...")
    for fila in range(2, max_row + 1):  # Comença des de la fila 2
        # Defineix les cel·les de dades per a aquesta fila
        cella_estudiant_dades = f"{col_estudiant}{fila}"
        cella_titol_dades     = f"{col_titol}{fila}"

        # Llegeix les dades amb protecció d'errors
        try:
            estudiant = str(ws_dades.range(cella_estudiant_dades).value).strip() if ws_dades.range(cella_estudiant_dades).value else ""
            titol = str(ws_dades.range(cella_titol_dades).value).strip() if ws_dades.range(cella_titol_dades).value else ""
        except Exception as e:
            print(f"⚠ Error llegint dades de la fila {fila}: {e}")
            print("⚠ Intentant reconnectar...")
            continue

        # Salta files buides

        if not estudiant:
            continue
        
        print(f"Processant: {estudiant}")            
        subdirectori = estudiant.replace(" ", "_")

        # Crea el subdirectori si no existeix
        if not os.path.exists(subdirectori):
            continue


        nom_fiter_excel = f"{os.path.splitext(plantilla_comite)[0]}_{subdirectori}.xlsx"
        nom_fiter_excel = os.path.join(subdirectori, nom_fiter_excel)

        nom_fitxer_pdf = f"Committee_rubric_{subdirectori}.pdf"
        nom_fitxer_pdf = os.path.join(subdirectori, nom_fitxer_pdf)

        if GENERAR_PDF:
            # Excel a PDF
            pdf_file = excel_to_pdf(nom_fiter_excel, nom_fitxer_pdf)
        
        if ENVIAR_EMAIL:
            # Only try to send email if PDF was created successfully

            if pdf_file:
                if not os.path.exists(pdf_file):
                    print(f"❌ El fitxer PDF no existeix: {pdf_file}")
                    continue
                # Send email with error handling
                try:
                    params = {}
                    params['destinatari'] = "ramon@cvc.uab.es"
                    params['assumpte'] = "Fitxer PDF generat - Rúbrica TFM"
                    params['nom_estudiant'] = estudiant
                    params['titol'] = titol
                    params['rol'] = 'president'
                    params['fitxer_pdf'] = pdf_file
                    params['plantilla_email'] = plantilla_email_signar

                    # Use the title and role for email content
                    email_sent = envia_email(params)
                    if not email_sent:
                        print(f"⚠ No s'ha pogut enviar l'email per {estudiant}")
                except Exception as e:
                    print(f"❌ Error enviant email per {estudiant}: {e}")
            else:
                print(f"❌ No s'ha pogut crear el PDF per {estudiant}, saltant l'enviament d'email")
            
            #print(f"✓ Processat: {estudiant}")
            #print("-" * 50)

    # Close the data file
    try:
        wb_dades.close()
    except:
        pass
    print("✓ Processament completat")

except Exception as e:
    print(f"Error durant l'execució: {e}")
    # Try to close the data file even if there was an error
    try:
        wb_dades.close()
    except:
        pass
finally:
    # Close the Excel application
    app.quit()
